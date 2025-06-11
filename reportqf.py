from docx import Document
from docx.shared import Inches
import pandas as pd
from utils.utils import get_df_serie_temporal, agregar_columna_mes
from plotly_utils.boxplot import fig_boxplot

# Cargar datos
ruta_prec = r"E:\045-Diego Rengifo 2024\04-PRODUCTOS SERIES TEMPORALES\Precipitaciones\control de calidad\BSO\BSO_PREC_CC.csv"
serie = get_df_serie_temporal(ruta_prec)
serie_q1 = serie[serie['QF'] == 1]

# Crear un nuevo documento
doc = Document()
doc.add_heading('Reporte Estadístico', level=1)

# Introducción
doc.add_heading('Introducción', level=2)
doc.add_paragraph('Este reporte analiza los datos de precipitaciones para el año 2020.')

# Cargar visualizaciones
boxplot_fig = fig_boxplot(serie_q1)
boxplot_fig.write_image("boxplot.png")  # Guarda la imagen

# Agregar visualización al documento
doc.add_heading('Gráfico de Boxplot', level=2)
doc.add_picture("boxplot.png", width=Inches(5.0))  # Ajusta el tamaño según sea necesario

# Conclusiones
doc.add_heading('Conclusiones', level=2)
doc.add_paragraph('Las precipitaciones mostraron un comportamiento ...')

# Guardar documento
doc.save('reporte_estadistico.docx')
